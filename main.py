import os
import json
import tempfile
from datetime import datetime

import requests
import qrcode
import gspread

from flask import Flask, request
from google.oauth2.service_account import Credentials
from docx import Document
from docx.shared import Inches
from num2words import num2words


# =========================
# НАСТРОЙКИ
# =========================

BOT_TOKEN = os.getenv("BOT_TOKEN")
GOOGLE_SHEET_ID = os.getenv("GOOGLE_SHEET_ID")
GOOGLE_SERVICE_ACCOUNT_JSON = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")

WEBHOOK_SECRET = "asia_mx_secret"

if not BOT_TOKEN:
    raise RuntimeError("Не задан BOT_TOKEN в Render Environment")

app = Flask(__name__)

USER_STATE = {}


# =========================
# РЕКВИЗИТЫ ASIA MX
# =========================

AGENT_NAME = "ЭЙЖА ЭМ ЭКС ТРЕЙД ЛИМИТЕД (ASIA MX TRADE LIMITED)"
AGENT_SHORT = "ЭЙЖА ЭМ ЭКС ТРЕЙД ЛИМИТЕД"

AGENT_ADDRESS = "Unit 704c 7/f Block 3, Nan Fung Industrial City, 18 Tin Hau Road, Tuen Mun, Hong Kong"
AGENT_EMAIL = "general@asiamx.ltd"
AGENT_WEB = "asiamx.ltd"
AGENT_TEL = "+852-9129-2658"

AGENT_INN = "9909722497"
AGENT_KPP = "270087001"

AGENT_ACCOUNT_PRINT = "40807 810 3 5071 0000002"
AGENT_ACCOUNT_QR = "40807810350710000002"

AGENT_BANK = "ДАЛЬНЕВОСТОЧНЫЙ БАНК ПАО СБЕРБАНК"
AGENT_BIK = "040813608"

AGENT_CORR_PRINT = "30101 810 6 0000 0000608"
AGENT_CORR_QR = "30101810600000000608"

BANK_INN = "7707083893"
BANK_KPP = "254002002"


# =========================
# TELEGRAM API
# =========================

def tg(method, payload=None, files=None):
    url = f"https://api.telegram.org/bot{BOT_TOKEN}/{method}"
    r = requests.post(url, data=payload, files=files, timeout=60)
    print(method, r.status_code, r.text[:500])
    return r


def send_message(chat_id, text, keyboard=None):
    payload = {
        "chat_id": chat_id,
        "text": text,
    }

    if keyboard:
        payload["reply_markup"] = json.dumps(keyboard, ensure_ascii=False)

    return tg("sendMessage", payload)


def send_document(chat_id, file_path, filename, caption=None):
    with open(file_path, "rb") as f:
        files = {
            "document": (filename, f),
        }
        payload = {
            "chat_id": chat_id,
        }
        if caption:
            payload["caption"] = caption

        return tg("sendDocument", payload, files)


def main_menu():
    return {
        "keyboard": [
            [{"text": "📄 Выставить счет"}],
            [{"text": "📑 Договор"}, {"text": "📎 Приложение"}],
            [{"text": "❌ Отмена"}],
        ],
        "resize_keyboard": True,
    }


def identifier_keyboard():
    return {
        "keyboard": [
            [{"text": "VIN"}, {"text": "Chassis"}],
            [{"text": "❌ Отмена"}],
        ],
        "resize_keyboard": True,
    }


def after_invoice_keyboard():
    return {
        "keyboard": [
            [{"text": "📑 Сформировать договор"}],
            [{"text": "📎 Сформировать приложение"}],
            [{"text": "🏠 В меню"}],
        ],
        "resize_keyboard": True,
    }


# =========================
# GOOGLE SHEETS
# =========================

def get_sheet():
    if not GOOGLE_SHEET_ID or not GOOGLE_SERVICE_ACCOUNT_JSON:
        return None

    info = json.loads(GOOGLE_SERVICE_ACCOUNT_JSON)

    scopes = [
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive",
    ]

    creds = Credentials.from_service_account_info(info, scopes=scopes)
    client = gspread.authorize(creds)

    return client.open_by_key(GOOGLE_SHEET_ID)


def get_next_invoice_number(data):
    sheet = get_sheet()

    if sheet is None:
        return datetime.now().strftime("%H%M%S")

    try:
        ws = sheet.worksheet("СЧЕТА")
    except Exception:
        ws = sheet.add_worksheet(title="СЧЕТА", rows=1000, cols=20)
        ws.append_row([
            "invoice_number",
            "created_at",
            "buyer_full_name",
            "contract_number",
            "contract_date",
            "foreign_invoice_number",
            "car_name",
            "identifier_type",
            "identifier_value",
            "amount",
        ])

    values = ws.get_all_values()

    last_number = 14252

    for row in values[1:]:
        if row and str(row[0]).isdigit():
            last_number = max(last_number, int(row[0]))

    new_number = last_number + 1

    ws.append_row([
        str(new_number),
        datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
        data.get("buyer_full_name", ""),
        data.get("contract_number", ""),
        data.get("contract_date", ""),
        data.get("foreign_invoice_number", ""),
        data.get("car_name", ""),
        data.get("identifier_type", ""),
        data.get("identifier_value", ""),
        data.get("amount", ""),
    ])

    return str(new_number)


# =========================
# ДЕНЬГИ / QR
# =========================

def normalize_amount(amount):
    return float(str(amount).replace(" ", "").replace(",", "."))


def format_amount(amount):
    value = normalize_amount(amount)
    return f"{value:,.2f}".replace(",", " ")


def amount_words(amount):
    value = int(normalize_amount(amount))
    text = num2words(value, lang="ru")
    return f"{text.capitalize()} рублей 00 копеек"


def amount_kopecks(amount):
    return str(int(round(normalize_amount(amount) * 100)))


def create_qr(data):
    purpose = (
        f"Оплата по счету №{data['invoice_number']} "
        f"по договору {data['contract_number']} от {data['contract_date']}"
    )

    qr_text = (
        f"ST00012|"
        f"Name={AGENT_SHORT}|"
        f"PersonalAcc={AGENT_ACCOUNT_QR}|"
        f"BankName={AGENT_BANK}|"
        f"BIC={AGENT_BIK}|"
        f"CorrespAcc={AGENT_CORR_QR}|"
        f"PayeeINN={AGENT_INN}|"
        f"KPP={AGENT_KPP}|"
        f"Purpose={purpose}|"
        f"Sum={amount_kopecks(data['amount'])}"
    )

    img = qrcode.make(qr_text)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    img.save(tmp.name)

    return tmp.name


# =========================
# ДОКУМЕНТЫ
# =========================

def create_invoice_docx(data):
    doc = Document()

    doc.add_paragraph(AGENT_BANK)
    doc.add_paragraph(f"Банк получателя")
    doc.add_paragraph(f"БИК {AGENT_BIK}")
    doc.add_paragraph(f"Корсчёт {AGENT_CORR_PRINT}")
    doc.add_paragraph(f"ИНН {AGENT_INN}    КПП {AGENT_KPP}")
    doc.add_paragraph(f"Сч. № {AGENT_ACCOUNT_PRINT}")
    doc.add_paragraph(f"Получатель: {AGENT_SHORT}")

    doc.add_heading(
        f"Счет на оплату № {data['invoice_number']} от {data['invoice_date']}",
        level=1,
    )

    doc.add_paragraph(
        f"Поставщик (Исполнитель): {AGENT_NAME}, ИНН {AGENT_INN}, КПП {AGENT_KPP}, "
        f"{AGENT_ADDRESS}"
    )

    doc.add_paragraph(
        f"Покупатель (Заказчик): {data['buyer_full_name']}"
    )

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    headers = ["№", "Товары (работы, услуги)", "Кол-во", "Ед.", "Цена", "Сумма"]

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    service_text = (
        f"Оплата по договору {data['contract_number']} от {data['contract_date']}, "
        f"инвойс {data['foreign_invoice_number']}, "
        f"за {data['car_name']}, "
        f"{data['identifier_type']} {data['identifier_value']}"
    )

    row = table.add_row().cells
    row[0].text = "1"
    row[1].text = service_text
    row[2].text = "1"
    row[3].text = "Шт"
    row[4].text = format_amount(data["amount"])
    row[5].text = format_amount(data["amount"])

    doc.add_paragraph("")
    doc.add_paragraph(f"Итого: {format_amount(data['amount'])}")
    doc.add_paragraph("НДС не облагается: -")
    doc.add_paragraph(f"Всего к оплате: {format_amount(data['amount'])}")

    doc.add_paragraph(
        f"Всего наименований 1, на сумму {format_amount(data['amount'])} руб."
    )
    doc.add_paragraph(amount_words(data["amount"]))

    qr_path = create_qr(data)
    doc.add_paragraph("")
    doc.add_paragraph("QR-код для оплаты:")
    doc.add_picture(qr_path, width=Inches(1.7))

    path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(path)

    return path


def create_contract_docx(data):
    doc = Document()

    doc.add_paragraph("ASIA MX TRADE LIMITED")
    doc.add_paragraph(AGENT_ADDRESS)
    doc.add_paragraph(f"Email: {AGENT_EMAIL}")
    doc.add_paragraph(f"Web: {AGENT_WEB}")
    doc.add_paragraph(f"Tel: {AGENT_TEL}")

    doc.add_heading(
        f"АГЕНТСКИЙ ДОГОВОР № {data['contract_number']}",
        level=1,
    )
    doc.add_paragraph("по организации оплаты автомобилей")
    doc.add_paragraph(f"г. Владивосток    {data['contract_date']}")

    doc.add_paragraph(
        f"{data['buyer_full_name']}, именуемый(ая) в дальнейшем «Принципал», "
        f"с одной стороны, и {AGENT_NAME}, зарегистрированная на территории САР Гонконг, "
        f"{AGENT_ADDRESS}, ИНН {AGENT_INN}, КПП {AGENT_KPP}, "
        f"в лице Антона Фалкона Пичардо, действующего на основании Устава, "
        f"именуемая в дальнейшем «Агент», с другой стороны, совместно именуемые "
        f"«Стороны», заключили настоящий договор о нижеследующем:"
    )

    doc.add_heading("1. Предмет договора", level=2)
    doc.add_paragraph(
        "1.1. Принципал поручает, а Агент принимает на себя обязательства по организации "
        "и проведению оплаты автомобилей, приобретаемых Принципалом, в соответствии "
        "с условиями, согласованными между Принципалом и соответствующими поставщиками."
    )
    doc.add_paragraph(
        "1.2. Агент осуществляет исключительно функции по приёму денежных средств "
        "от Принципала и их перечислению на счета поставщиков, указанных Принципалом."
    )

    doc.add_heading("2. Обязанности Агента", level=2)
    doc.add_paragraph(
        "2.1. Получать от Принципала письменные или электронные поручения "
        "с указанием реквизитов поставщиков и суммы платежа."
    )
    doc.add_paragraph(
        "2.2. Перечислять денежные средства поставщикам в сроки, согласованные с Принципалом."
    )
    doc.add_paragraph(
        "2.3. Предоставлять Принципалу подтверждающие документы об отправке платежей."
    )

    doc.add_heading("3. Обязанности Принципала", level=2)
    doc.add_paragraph(
        "3.1. Своевременно предоставлять Агенту необходимые реквизиты поставщика и сумму платежа."
    )
    doc.add_paragraph(
        "3.2. Обеспечить поступление денежных средств на расчётный счёт Агента "
        "до даты перечисления средств поставщику."
    )

    doc.add_heading("4. Вознаграждение и порядок расчётов", level=2)
    doc.add_paragraph(
        "4.1. Вознаграждение Агента составляет 1% от суммы каждого платежа по курсу ЦБ."
    )
    doc.add_paragraph(
        "4.2. Оплата вознаграждения производится путём безналичного перечисления "
        "на расчётный счёт Агента в момент перевода денежных средств."
    )

    doc.add_heading("5. Банковские реквизиты сторон", level=2)
    doc.add_paragraph(
        f"Агент:\n"
        f"ASIA MX TRADE LIMITED\n"
        f"{AGENT_ADDRESS}\n"
        f"Email: {AGENT_EMAIL}\n"
        f"Web: {AGENT_WEB}\n"
        f"Tel: {AGENT_TEL}\n\n"
        f"Банковские реквизиты (Россия):\n"
        f"Наименование: {AGENT_NAME}\n"
        f"ИНН: {AGENT_INN}, КПП: {AGENT_KPP}\n"
        f"Р/с: {AGENT_ACCOUNT_PRINT}\n"
        f"Банк: {AGENT_BANK}\n"
        f"БИК: {AGENT_BIK}\n"
        f"Корсчёт: {AGENT_CORR_PRINT}\n"
        f"ИНН банка: {BANK_INN}, КПП банка: {BANK_KPP}"
    )

    doc.add_heading("6. Ответственность сторон", level=2)
    doc.add_paragraph(
        "6.1. Стороны несут ответственность за ненадлежащее исполнение своих обязательств "
        "в соответствии с законодательством."
    )
    doc.add_paragraph(
        "6.2. Агент освобождается от ответственности за любые убытки, вызванные "
        "действиями/бездействием поставщиков."
    )

    doc.add_heading("7. Форс-мажор", level=2)
    doc.add_paragraph(
        "Стороны освобождаются от ответственности за частичное или полное неисполнение "
        "обязательств по настоящему договору, если оно явилось следствием обстоятельств "
        "непреодолимой силы."
    )

    doc.add_heading("8. Срок действия договора", level=2)
    doc.add_paragraph(
        "8.1. Настоящий договор вступает в силу с момента подписания и действует "
        "до «31» декабря 2026 г."
    )

    doc.add_heading("9. Заключительные положения", level=2)
    doc.add_paragraph(
        "9.1. Все споры и разногласия решаются путём переговоров, а при недостижении "
        "согласия — разрешаются на основании российского законодательства."
    )

    doc.add_paragraph("")
    doc.add_paragraph("Подписи сторон:")
    doc.add_paragraph(f"Агент: _____________________ /Антон Фалкон Пичардо/")
    doc.add_paragraph(f"Принципал: _____________________ /{data['buyer_full_name']}/")

    path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(path)

    return path


def create_appendix_docx(data):
    doc = Document()

    doc.add_heading(
        f"Приложение №1 к Агентскому договору № {data['contract_number']} от {data['contract_date']}",
        level=1,
    )

    doc.add_heading(
        f"Поручение Агенту №1 от {data['invoice_date']}",
        level=2,
    )

    doc.add_paragraph(
        f"{data['buyer_full_name']} (Принципал) поручает {AGENT_NAME} (Агент) "
        f"обеспечить своевременную оплату в пользу третьего лица (Контрагента) "
        f"от имени Принципала по договору поставки."
    )

    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"

    after_fee = normalize_amount(data["amount"]) * 0.99

    rows = [
        ("Валюта-1", "Российский рубль"),
        ("Сумма платежа, Валюта-1", f"{format_amount(data['amount'])} руб."),
        (
            "Банковские реквизиты Агента",
            f"{AGENT_NAME}\n"
            f"ИНН: {AGENT_INN}, КПП: {AGENT_KPP}\n"
            f"Р/с: {AGENT_ACCOUNT_PRINT}\n"
            f"Банк: {AGENT_BANK}\n"
            f"БИК: {AGENT_BIK}\n"
            f"Корсчёт: {AGENT_CORR_PRINT}\n"
            f"ИНН банка: {BANK_INN}, КПП банка: {BANK_KPP}",
        ),
        ("Вознаграждение Агента", "1%"),
        ("Валюта-2", "Российский рубль"),
        ("Сумма после удержания вознаграждения", f"{format_amount(after_fee)} руб."),
        ("Контрагент", "Заполняется дополнительно"),
        ("Другие условия", ""),
    ]

    for i, row_data in enumerate(rows):
        table.rows[i].cells[0].text = row_data[0]
        table.rows[i].cells[1].text = row_data[1]

    doc.add_paragraph("")
    doc.add_paragraph(f"Агент:\n{AGENT_NAME}\n_____________________ /Антон Фалкон Пичардо/")
    doc.add_paragraph("")
    doc.add_paragraph(f"Принципал:\n_____________________ /{data['buyer_full_name']}/")

    path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(path)

    return path


# =========================
# ЛОГИКА БОТА
# =========================

def reset(chat_id):
    USER_STATE.pop(str(chat_id), None)


def start_invoice(chat_id):
    USER_STATE[str(chat_id)] = {
        "mode": "invoice",
        "step": "buyer_full_name",
        "data": {},
    }

    send_message(chat_id, "Введите ФИО покупателя:")


def process_invoice_step(chat_id, text):
    state = USER_STATE.get(str(chat_id))

    if not state:
        send_message(chat_id, "Выберите действие:", main_menu())
        return

    step = state["step"]
    data = state["data"]

    if step == "buyer_full_name":
        data["buyer_full_name"] = text
        state["step"] = "contract_number"
        send_message(chat_id, "Введите номер договора:")
        return

    if step == "contract_number":
        data["contract_number"] = text
        state["step"] = "contract_date"
        send_message(chat_id, "Введите дату договора:")
        return

    if step == "contract_date":
        data["contract_date"] = text
        state["step"] = "foreign_invoice_number"
        send_message(chat_id, "Введите номер инвойса:")
        return

    if step == "foreign_invoice_number":
        data["foreign_invoice_number"] = text
        state["step"] = "car_name"
        send_message(chat_id, "Введите наименование авто / марка, модель:")
        return

    if step == "car_name":
        data["car_name"] = text
        state["step"] = "identifier_type"
        send_message(chat_id, "Выберите тип идентификатора:", identifier_keyboard())
        return

    if step == "identifier_type":
        if text not in ["VIN", "Chassis"]:
            send_message(chat_id, "Выберите VIN или Chassis:", identifier_keyboard())
            return

        data["identifier_type"] = text
        state["step"] = "identifier_value"
        send_message(chat_id, f"Введите значение {text}:")
        return

    if step == "identifier_value":
        data["identifier_value"] = text
        state["step"] = "amount"
        send_message(chat_id, "Введите сумму счета:")
        return

    if step == "amount":
        data["amount"] = text.replace(" ", "").replace(",", ".")
        data["invoice_date"] = datetime.now().strftime("%d.%m.%Y")
        data["invoice_number"] = get_next_invoice_number(data)

        preview = (
            "Проверьте данные:\n\n"
            f"ФИО: {data['buyer_full_name']}\n"
            f"Счет № {data['invoice_number']} от {data['invoice_date']}\n"
            f"Договор: {data['contract_number']} от {data['contract_date']}\n"
            f"Инвойс: {data['foreign_invoice_number']}\n"
            f"Авто: {data['car_name']}\n"
            f"{data['identifier_type']}: {data['identifier_value']}\n"
            f"Сумма: {format_amount(data['amount'])} руб.\n\n"
            f"Формирую счет..."
        )

        send_message(chat_id, preview)

        invoice_path = create_invoice_docx(data)

        send_document(
            chat_id,
            invoice_path,
            f"Счет_{data['invoice_number']}.docx",
            "Счет сформирован ✅",
        )

        send_message(
            chat_id,
            "Что формируем дальше?",
            after_invoice_keyboard(),
        )

        state["step"] = "done"
        return

    if step == "done":
        send_message(chat_id, "Документы уже сформированы. Выберите действие:", after_invoice_keyboard())
        return


def create_contract_from_last(chat_id):
    state = USER_STATE.get(str(chat_id))

    if not state or not state.get("data"):
        send_message(chat_id, "Сначала сформируйте счет, чтобы договор подтянул данные.")
        return

    data = state["data"]
    path = create_contract_docx(data)

    send_document(
        chat_id,
        path,
        f"Договор_{data['contract_number']}.docx",
        "Договор сформирован ✅",
    )


def create_appendix_from_last(chat_id):
    state = USER_STATE.get(str(chat_id))

    if not state or not state.get("data"):
        send_message(chat_id, "Сначала сформируйте счет, чтобы приложение подтянуло данные.")
        return

    data = state["data"]
    path = create_appendix_docx(data)

    send_document(
        chat_id,
        path,
        f"Приложение_{data['contract_number']}.docx",
        "Приложение сформировано ✅",
    )


def handle_message(message):
    chat_id = message["chat"]["id"]
    text = message.get("text", "").strip()

    if not text:
        send_message(chat_id, "Я пока работаю только с текстом.")
        return

    if text == "/start":
        reset(chat_id)
        send_message(
            chat_id,
            "Бот документов ASIA MX запущен ✅\n\nВыберите действие:",
            main_menu(),
        )
        return

    if text in ["🏠 В меню", "❌ Отмена"]:
        reset(chat_id)
        send_message(chat_id, "Выберите действие:", main_menu())
        return

    if text == "📄 Выставить счет":
        start_invoice(chat_id)
        return

    if text in ["📑 Договор", "📑 Сформировать договор"]:
        create_contract_from_last(chat_id)
        return

    if text in ["📎 Приложение", "📎 Сформировать приложение"]:
        create_appendix_from_last(chat_id)
        return

    state = USER_STATE.get(str(chat_id))

    if state and state.get("mode") == "invoice":
        process_invoice_step(chat_id, text)
        return

    send_message(chat_id, "Выберите действие:", main_menu())


# =========================
# FLASK ROUTES
# =========================

@app.route("/")
def home():
    return "ASIA MX Docs Bot is running ✅", 200


@app.route("/test")
def test():
    return "TEST OK", 200


@app.route(f"/webhook/{WEBHOOK_SECRET}", methods=["GET", "POST"])
def webhook():
    if request.method == "GET":
        return "Webhook endpoint is alive", 200

    update = request.get_json(force=True)
    print(update)

    try:
        if "message" in update:
            handle_message(update["message"])
    except Exception as e:
        print("BOT ERROR:", repr(e))

        try:
            chat_id = update["message"]["chat"]["id"]
            send_message(chat_id, f"Ошибка: {e}")
        except Exception:
            pass

    return "ok", 200


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)
